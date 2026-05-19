#! python3
# customInvitations.py - Create Custom Invitations from text file Guest List

import docx
from pathlib import Path

guestFile = open(Path.cwd() / 'guests.txt')

guestList = guestFile.readlines()

doc = docx.Document()

iterator = 0

for guestNum, guest in enumerate(guestList):
    doc.add_paragraph('It would be a pleasure to have the company of')
    doc.paragraphs[iterator].style = 'Subtitle'
    doc.paragraphs[iterator].runs[0].bold = True
    iterator = iterator + 1
    doc.add_paragraph(guest)
    doc.paragraphs[iterator].style = 'Caption'
    doc.paragraphs[iterator].runs[0].bold = True
    iterator = iterator + 1
    doc.add_paragraph('at 11010 Memory Lane of the Evening of')
    doc.paragraphs[iterator].style = 'Subtitle'
    doc.paragraphs[iterator].runs[0].bold = True
    iterator = iterator + 1
    doc.add_paragraph('April 1st')
    doc.paragraphs[iterator].style = 'Caption'
    doc.paragraphs[iterator].runs[0].bold = False
    iterator = iterator + 1
    doc.add_paragraph("at 7 o'clock")
    doc.paragraphs[iterator].style = 'Subtitle'
    doc.paragraphs[iterator].runs[0].bold = True
    doc.paragraphs[iterator].runs[0].add_break(docx.enum.text.WD_BREAK.PAGE)
    iterator = iterator + 1

doc.save('Invitations.docx')
